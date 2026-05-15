
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Style settings
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Title
    title = doc.add_heading('제약 규제 과학 자산 상세 인벤토리 보고서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('작성일: 2026년 05월 12일\n작성자: Antigravity AI Assistant')
    doc.add_page_break()

    # Table of Contents (Simulated)
    doc.add_heading('1. 개요 (Executive Summary)', level=1)
    doc.add_paragraph(
        "본 보고서는 사용자님의 로컬 시스템 내 산재된 제약 관련 전문 문서를 온톨로지(Ontology)적 관점에서 재분류하고 상세 내용을 정리한 것입니다. "
        "확인된 자산은 크게 식약처 대응(RA), 품질 및 시험법(CMC), 독성 평가(Safety), 글로벌 인허가 전략의 4가지 핵심 도메인으로 구분됩니다."
    )

    # RA Domain
    doc.add_heading('2. 규제 대응 및 품목 허가 (RA Domain)', level=1)
    
    # 2.1 Exiglue
    doc.add_heading('2.1. 엑시글루엠서방정 식약처보완답변서', level=2)
    doc.add_paragraph(
        "• 상세 내용: 서방성 제제의 용출 양상 및 유연물질 관리에 대한 식약처 지적 사항에 대한 논리적 방어 자료.\n"
        "• 세부 목차:\n"
        "  1) 보완 요구 사항 개요\n"
        "  2) 용출 시험 결과 및 타당성 (Dissolution Profile)\n"
        "  3) 유연물질 관리 기준 및 분석 데이터 (Impurity Control)\n"
        "  4) 분석법 밸리데이션(MV) 보완 결과\n"
        "  5) 결론 및 향후 관리 계획"
    )

    # 2.2 Minurin
    doc.add_heading('2.2. 미뉴린정(2함량) 기시 2차 보완시행문', level=2)
    doc.add_paragraph(
        "• 상세 내용: 데스모프레신 제제의 기준 및 시험방법(기시법)에 대한 식약처의 2차 보완 요구 사항.\n"
        "• 세부 목차:\n"
        "  1) 1차 보완 답변에 대한 식약처 추가 검토 의견\n"
        "  2) 원료 및 완제의 안정성 시험(Stability) 데이터 요구\n"
        "  3) 유연물질 규격 설정 근거 재검토\n"
        "  4) 최종 허가 예정일 및 행정 사항"
    )

    # Safety Domain
    doc.add_heading('3. 비임상 독성 및 안전성 (Safety Domain)', level=1)
    
    # 3.1 Naltrexone
    doc.add_heading('3.1. Naltrexone ER 주사제 AI 독성 CTD 리포트', level=2)
    doc.add_paragraph(
        "• 상세 내용: ICH M7 가이드라인에 따른 AI 기반 유전독성 예측 및 CTD 보고서.\n"
        "• 세부 목차:\n"
        "  1) Introduction (약물 개요 및 투여 경로)\n"
        "  2) In Silico Methodology (QSAR 모델 및 알고리즘 설명)\n"
        "  3) Structural Alert Analysis (독성 발현 구조 분석)\n"
        "  4) ICH M7 Classification (불순물 분류 1-5단계)\n"
        "  5) Expert Review & Conclusion (전문가 검토 및 결론)"
    )

    # CMC Domain
    doc.add_heading('4. 품질 및 기술문서 (CMC Domain)', level=1)
    
    # 4.1 CKD-830
    doc.add_heading('4.1. CKD-830 기준 및 시험방법 (Specification)', level=2)
    doc.add_paragraph(
        "• 상세 내용: 개발 중인 신약/개량신약의 품질 규격 설정 근거 자료.\n"
        "• 세부 목차:\n"
        "  1) Drug Substance (원료의약품) 성상 및 확인 시험\n"
        "  2) 순도 시험 (유연물질, 잔류용매)\n"
        "  3) 함량 시험 및 분석법 설명\n"
        "  4) Drug Product (완제의약품) 제제 특성 분석"
    )

    # 4.2 Sokchung
    doc.add_heading('4.2. 속청케어액/큐액 시험방법 밸리데이션(MV) 보고서', level=2)
    doc.add_paragraph(
        "• 상세 내용: 액제 제형의 분석법 타당성을 입증한 실험 데이터.\n"
        "• 세부 목차:\n"
        "  1) 특이성 (Specificity)\n"
        "  2) 직선성 (Linearity) 및 범위\n"
        "  3) 정확성 (Accuracy) 및 정밀성 (Precision)\n"
        "  4) 검출한계(LOD) 및 정량한계(LOQ)"
    )

    # Strategy Domain
    doc.add_heading('5. 글로벌 인허가 및 전략 (Strategy Domain)', level=1)
    doc.add_paragraph(
        "• 미국 FDA 의료기기/의약품 허가 제도 요약\n"
        "• Merck 및 종근당(CKD) 파이프라인 분석 보고서\n"
        "• 국내외 허가 제도 비교 및 대응 시나리오"
    )

    # Footer
    doc.add_paragraph('\n[보고서 종료]')

    # Save
    output_path = '/Users/leeyoung-nam/Desktop/Antigravity Project/Regulatory_Asset_Inventory_Report.docx'
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    path = create_report()
    print(f"Report created at: {path}")
